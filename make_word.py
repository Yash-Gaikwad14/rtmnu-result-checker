import docx
import sys
import re

def create_docx_from_md(md_path, docx_path):
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    for line in content.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. '):
            doc.add_paragraph(line[3:], style='List Number')
        elif line.strip() == '---':
            doc.add_paragraph('_' * 30)
        elif line.strip():
            # Basic bold parsing for paragraphs
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

    doc.save(docx_path)
    print("Saved to " + docx_path)

if __name__ == "__main__":
    create_docx_from_md(r"C:\Users\Yashg\.gemini\antigravity\brain\76ef3d9d-b1d0-4453-84ef-fcec5932b912\walkthrough.md", r"c:\Users\Yashg\Downloads\New folder (2)\RTMNU_Project_Walkthrough.docx")
